"""Generate ATS-optimized resume for Remedora Full-Stack Laravel Engineer role."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "/workspace/resumes/Erik_Gratz_Resume_Remedora.docx"


def add_heading(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(8)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Name & contact
    name = doc.add_paragraph()
    name_run = name.add_run("Erik Gratz")
    name_run.bold = True
    name_run.font.size = Pt(16)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    contact_run = contact.add_run(
        "Remote | 937-218-2673 | erikgratz110@gmail.com | "
        "https://www.linkedin.com/in/erik-gratz-126ba410b/ | https://erikgratz.com"
    )
    contact_run.font.size = Pt(9)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)

    # Summary
    add_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Full-stack Laravel engineer with 6+ years of production PHP experience building "
        "scalable APIs, queue-driven integrations, and cloud-hosted platforms in wellness and "
        "fintech. Owned backend development from MVP through acquisition at an early-stage "
        "startup, scaling systems 40x while improving query performance, data integrity, and "
        "deployment workflows. Experienced with Laravel queues (Horizon, RabbitMQ), third-party "
        "API integration (including maintaining SDKs for poorly documented providers), Docker, "
        "CI/CD, PHPUnit testing, access control, and secure handling of sensitive user data. "
        "Strong Eloquent and SQL skills with a track record of live-data migrations, schema "
        "refactors, and performance optimization without downtime."
    )
    p = doc.add_paragraph(summary)
    for run in p.runs:
        run.font.size = Pt(10)

    # Skills - keyword dense for ATS
    add_heading(doc, "TECHNICAL SKILLS")
    skills = [
        "Languages & Frameworks: PHP 8.x, Laravel, Lumen, Eloquent, Livewire, FilamentPHP, "
        "JavaScript, Vue.js, REST APIs, API design (consuming and exposing)",
        "Data & Performance: MySQL, MongoDB, SQL query optimization, schema design, "
        "database migrations on live production data, ElasticSearch",
        "Queues & Integrations: Laravel Horizon, queued background jobs, RabbitMQ, webhook "
        "and cron workflows, third-party API integration, retry and reconciliation patterns",
        "Testing & Quality: PHPUnit, unit testing, code coverage, GitHub Actions CI/CD",
        "Infrastructure & DevOps: Docker, Laravel Forge, Sail, Google Cloud Platform, AWS, "
        "DigitalOcean, zero-downtime deployments",
        "Security & Compliance: access control, permissions, API key management, XSS and "
        "vulnerability remediation, security-conscious development in regulated-adjacent "
        "wellness and financial products",
        "Payments & Financial APIs: Plaid API integration and SDK maintenance",
        "Laravel Ecosystem: Passport, Sanctum, Telescope, Pint, Flare",
    ]
    for skill in skills:
        add_bullet(doc, skill)

    # Experience
    add_heading(doc, "PROFESSIONAL EXPERIENCE")

    add_body(doc, "Senior Software Engineer, Backend | Pocketnest | Detroit, MI (Remote)", bold=True)
    add_body(doc, "May 2022 – Present")
    pocketnest = [
        "Led Laravel backend development from MVP through acquisition, scaling the platform "
        "to 40x user growth while supporting a 36.6x increase in company valuation at exit.",
        "Built and maintained production APIs with Laravel, including a client-facing analytics "
        "API with API key management, role-based permissions, and email invitation workflows.",
        "Owned Laravel Horizon queue infrastructure for background jobs, external service "
        "calls, and async processing across staging and production environments.",
        "Executed zero-downtime migration of the full platform from AWS to Google Cloud using "
        "Docker and Laravel Forge for deployment and server management.",
        "Optimized Eloquent queries and application routes, achieving up to 20x performance "
        "improvements in the first month through SQL tuning and architectural refactors.",
        "Refactored application core to eliminate data redundancy and database bloat, improving "
        "maintainability and query performance across the schema.",
        "Built an internal FilamentPHP configuration dashboard with database snapshots, "
        "analytics, and staging-to-production config transfer for faster iteration.",
        "Maintained a fork of the Plaid PHP SDK to support evolving third-party API "
        "requirements when the upstream library was abandoned.",
        "Technologies: PHP, Laravel, MySQL, Eloquent, Docker, FilamentPHP, Livewire, Horizon, "
        "Passport, Sanctum, Telescope, Forge, Sail, PHPUnit",
    ]
    for b in pocketnest:
        add_bullet(doc, b)

    add_body(doc, "Developer | Sonic Boom Wellness | San Diego, CA (Remote)", bold=True)
    add_body(doc, "February 2020 – May 2022")
    sonic = [
        "Developed and maintained APIs for an employee wellness platform in a regulated-adjacent "
        "healthcare context, with focus on scalability, data validation, and secure access.",
        "Migrated platform features from a legacy Zend application to a new Lumen API, "
        "modernizing the PHP stack and improving maintainability.",
        "Implemented extensive queued jobs, cron scheduling, and RabbitMQ-backed background "
        "processing for integrations with external systems and large data imports.",
        "Led a self-directed security initiative with staff to identify and remediate "
        "vulnerabilities across the stack, including XSS and input validation gaps.",
        "Technologies: PHP, Laravel, Lumen, Eloquent, MySQL, MongoDB, Vue 3, JavaScript, "
        "RabbitMQ, ElasticSearch, Rundeck, Docker",
    ]
    for b in sonic:
        add_bullet(doc, b)

    add_body(doc, "PHP Developer (Lead) | Internet Things | San Diego, CA", bold=True)
    add_body(doc, "April 2019 – February 2020")
    internet = [
        "Stepped into lead developer role and owned continued development of SimplySweeps.com, "
        "a lead-generation platform with automated third-party API integrations.",
        "Built and maintained an API-driven lead auction system with transparent reporting, "
        "connection debugging UI, and real-time integration with external partner APIs.",
        "Technologies: PHP, Laravel, Eloquent, MySQL, JavaScript, jQuery, Docker, JIRA",
    ]
    for b in internet:
        add_bullet(doc, b)

    # Open source
    add_heading(doc, "OPEN SOURCE & PROJECTS")
    projects = [
        "Plaid SDK for PHP (2024–Present): Maintains abandoned upstream SDK for ongoing Plaid "
        "API compatibility; demonstrates third-party integration and API versioning discipline.",
        "Code Coverage Summary GitHub Action (2025–Present): CI/CD tooling for targeted "
        "PHPUnit/cobertura coverage analysis on pull requests.",
        "Wix SDK for PHP (2023–Present): Custom PHP client for Wix website API integration.",
        "Personal portfolio site: Laravel, Livewire, MySQL, Tailwind CSS, Docker, Forge, "
        "DigitalOcean, with PHPUnit test suite.",
    ]
    for p in projects:
        add_bullet(doc, p)

    # Education
    add_heading(doc, "EDUCATION & CONTINUING EDUCATION")
    edu = [
        "Laracasts — Laravel-focused professional development (2020–Present)",
        "Treehouse — Full Stack Development (2019)",
        "Ohio University — Computer Science coursework (2008–2010)",
    ]
    for e in edu:
        add_bullet(doc, e)

    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")


if __name__ == "__main__":
    build()
